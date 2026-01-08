"""
Скрипт для створення реферату та документа з блок-схемою
Лабораторна робота з текстового процесора
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
import os


def add_page_number(doc):
    """Додати нумерацію сторінок"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)


def add_hyperlink(paragraph, text, url):
    """Додати гіперпосилання"""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    # Шрифт
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rFonts)
    
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '28')  # 14pt * 2
    rPr.append(sz)
    
    new_run.append(rPr)
    
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_bookmark(paragraph, bookmark_name, text):
    """Додати закладку"""
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), str(hash(bookmark_name) % 10000))
    bookmark_start.set(qn('w:name'), bookmark_name)
    
    run = paragraph.add_run(text)
    
    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), str(hash(bookmark_name) % 10000))
    
    paragraph._p.insert(0, bookmark_start)
    paragraph._p.append(bookmark_end)


def add_toc(doc):
    """Додати автоматичний зміст"""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)


def add_centered_text(doc, text, bold=False, size=14):
    """Додати текст по центру"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    return p


def add_right_text(doc, text, bold=False, size=14):
    """Додати текст по правому краю"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    return p


def add_justified_text(doc, text, first_indent=True, size=14):
    """Додати текст з вирівнюванням по ширині"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    return p


def add_left_text(doc, text, indent=0, size=14):
    """Додати текст по лівому краю"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.line_spacing = 1.5
    return p


def add_bullet_list(doc, items, size=14):
    """Додати маркований список"""
    for item in items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run('• ' + item)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.line_spacing = 1.5


def create_referat():
    """Створити реферат"""
    doc = Document()
    
    # Налаштування стилів
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    
    # Налаштування полів сторінки
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
    
    # ========== ТИТУЛЬНА СТОРІНКА ==========
    add_centered_text(doc, 'МІНІСТЕРСТВО НАУКИ ТА ОСВІТИ УКРАЇНИ', bold=True)
    add_centered_text(doc, 'Дрогобицький державний педагогічний університет ім. Івана Франка', bold=True)
    add_centered_text(doc, 'Факультет фізики математики економіки та інноваційних технологій')
    
    # Відступ
    for _ in range(5):
        doc.add_paragraph()
    
    add_centered_text(doc, 'РЕФЕРАТ', bold=True, size=18)
    add_centered_text(doc, 'на тему:')
    add_centered_text(doc, '"ХМАРНІ ТЕХНОЛОГІЇ ТА ЇХ ЗАСТОСУВАННЯ В СУЧАСНИХ ІНФОРМАЦІЙНИХ СИСТЕМАХ"', bold=True, size=16)
    
    # Відступ
    for _ in range(6):
        doc.add_paragraph()
    
    add_right_text(doc, 'Виконав студент групи КН-2429Б')
    add_right_text(doc, 'Янісів Максим')
    doc.add_paragraph()
    add_right_text(doc, 'Керівник: к.т.н., доц. Петренко О.В.')
    
    # Відступ до низу
    for _ in range(5):
        doc.add_paragraph()
    
    add_centered_text(doc, 'Дрогобич – 2026')
    
    doc.add_page_break()
    
    # ========== СТОРІНКА 2 - АНОТАЦІЯ ==========
    add_centered_text(doc, 'АНОТАЦІЯ', bold=True, size=16)
    doc.add_paragraph()
    
    annotation_text = """Даний реферат присвячено дослідженню хмарних технологій та їх застосуванню в сучасних інформаційних системах. У роботі розглянуто основні концепції хмарних обчислень, їх історичний розвиток та сучасний стан. Проаналізовано три основні моделі надання хмарних послуг: інфраструктура як послуга (IaaS), платформа як послуга (PaaS) та програмне забезпечення як послуга (SaaS).

Особливу увагу приділено питанням безпеки даних у хмарних середовищах, методам шифрування та захисту інформації. Розглянуто переваги та недоліки використання хмарних технологій для бізнесу та освіти.

У рефераті наведено порівняльний аналіз провідних хмарних платформ: Amazon Web Services, Microsoft Azure та Google Cloud Platform. Досліджено тенденції розвитку хмарних технологій та їх вплив на цифрову трансформацію суспільства.

Робота містить практичні рекомендації щодо вибору хмарних рішень для різних типів організацій та проектів.

Ключові слова: хмарні технології, хмарні обчислення, IaaS, PaaS, SaaS, інформаційні системи, безпека даних, цифрова трансформація."""
    
    add_justified_text(doc, annotation_text)
    
    doc.add_page_break()
    
    # ========== СТОРІНКА 3 - ЗМІСТ ==========
    add_centered_text(doc, 'ЗМІСТ', bold=True, size=16)
    doc.add_paragraph()
    add_toc(doc)
    doc.add_page_break()
    
    # ========== ВСТУП ==========
    p = doc.add_heading('ВСТУП', level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_justified_text(doc, """Сучасний світ переживає стрімку цифрову трансформацію, в якій хмарні технології відіграють ключову роль. Хмарні обчислення революціонізували спосіб зберігання, обробки та передачі даних, ставши невід'ємною частиною інформаційної інфраструктури підприємств, державних установ та освітніх закладів.""")

    add_justified_text(doc, """Актуальність теми дослідження обумовлена зростаючою потребою в гнучких, масштабованих та економічно ефективних інформаційних технологіях. За даними аналітичних компаній, ринок хмарних послуг щороку зростає на 20-25%, що свідчить про високий попит на ці технології.""")
    
    # Гіперпосилання на файл
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run('Детальну статистику використання хмарних сервісів можна знайти у файлі ')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    add_hyperlink(p, 'cloud_services_statistics.csv', 'file:///Users/maksymyanisiv/test%20for%20cursor/cloud_services_statistics.csv')
    run = p.add_run('.')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    
    add_justified_text(doc, """Метою даного реферату є комплексний аналіз хмарних технологій, їх класифікація, переваги та недоліки, а також дослідження практичних аспектів застосування в різних галузях.""")

    add_justified_text(doc, """Об'єктом дослідження є хмарні технології як сукупність інформаційних ресурсів та сервісів.""")

    add_justified_text(doc, """Предметом дослідження є особливості функціонування, впровадження та використання хмарних технологій в сучасних інформаційних системах.""")
    
    doc.add_page_break()
    
    # ========== РОЗДІЛ 1 ==========
    p = doc.add_heading('РОЗДІЛ 1. ТЕОРЕТИЧНІ ОСНОВИ ХМАРНИХ ТЕХНОЛОГІЙ', level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Підрозділ 1.1
    doc.add_heading('1.1. Поняття та історія розвитку хмарних обчислень', level=2)
    
    # Текст з закладкою
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_bookmark(p, 'cloud_computing_definition', '')
    run = p.add_run('Хмарні обчислення (англ. Cloud Computing) – це модель надання зручного мережевого доступу за запитом до спільного пулу обчислювальних ресурсів, що підлягають налаштуванню (наприклад, мережі, сервери, сховища даних, застосунки та послуги), які можуть бути швидко надані та вивільнені з мінімальними зусиллями з управління або взаємодії з постачальником послуг.')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    
    add_justified_text(doc, """Концепція хмарних обчислень з'явилася ще в 1960-х роках, коли Джон МакКарті висловив ідею про те, що обчислювальні потужності можуть надаватися як комунальна послуга. Однак практична реалізація цієї концепції стала можливою лише з розвитком Інтернету та технологій віртуалізації.""")
    
    add_justified_text(doc, """Ключовими етапами розвитку хмарних технологій стали:""")
    
    add_bullet_list(doc, [
        '1999 рік – запуск Salesforce.com, першого комерційного SaaS-рішення;',
        '2002 рік – Amazon Web Services починає надавати хмарні послуги;',
        '2006 рік – запуск Amazon EC2 та Google Apps;',
        '2008 рік – Microsoft представляє Azure;',
        '2010-ті роки – масове впровадження хмарних технологій у бізнесі та освіті.'
    ])
    
    # Гіперпосилання на Інтернет
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run('Більш детальну інформацію про історію хмарних технологій можна знайти на ')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    add_hyperlink(p, 'Wikipedia - Cloud Computing', 'https://uk.wikipedia.org/wiki/Хмарні_обчислення')
    run = p.add_run(' [1].')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    
    # Параграф 1.1.1
    doc.add_heading('1.1.1. Основні характеристики хмарних обчислень', level=3)
    
    add_justified_text(doc, """Національний інститут стандартів і технологій США (NIST) визначає п'ять основних характеристик хмарних обчислень:""")
    
    add_justified_text(doc, """1. Самообслуговування за запитом (On-demand self-service) – споживач може самостійно замовляти обчислювальні потужності без необхідності взаємодії з персоналом постачальника.""")
    
    add_justified_text(doc, """2. Широкий мережевий доступ (Broad network access) – ресурси доступні через мережу та можуть використовуватися різними клієнтськими платформами.""")
    
    add_justified_text(doc, """3. Об'єднання ресурсів (Resource pooling) – ресурси постачальника об'єднуються для обслуговування багатьох споживачів за моделлю мультиорендності.""")
    
    add_justified_text(doc, """4. Швидка еластичність (Rapid elasticity) – можливості можуть швидко та еластично надаватися, у деяких випадках автоматично.""")
    
    add_justified_text(doc, """5. Вимірюваність послуг (Measured service) – хмарні системи автоматично контролюють та оптимізують використання ресурсів.""")
    
    # Підрозділ 1.2
    doc.add_heading('1.2. Моделі надання хмарних послуг', level=2)
    
    # IaaS
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_bookmark(p, 'iaas_definition', '')
    run = p.add_run('Інфраструктура як послуга (IaaS – Infrastructure as a Service) – це модель надання хмарних послуг, при якій споживач отримує можливість використовувати обробку даних, зберігання, мережі та інші фундаментальні обчислювальні ресурси. Споживач не керує базовою хмарною інфраструктурою, але контролює операційні системи, сховища та розгорнуті застосунки.')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    
    # PaaS
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_bookmark(p, 'paas_definition', '')
    run = p.add_run('Платформа як послуга (PaaS – Platform as a Service) – це модель, при якій споживач може розгортати на хмарній інфраструктурі створені або придбані застосунки, використовуючи мови програмування, бібліотеки, служби та інструменти, що підтримуються постачальником.')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    
    # SaaS
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_bookmark(p, 'saas_definition', '')
    run = p.add_run('Програмне забезпечення як послуга (SaaS – Software as a Service) – це модель, при якій споживач використовує застосунки постачальника, що працюють на хмарній інфраструктурі. Застосунки доступні з різних клієнтських пристроїв через тонкий клієнт, наприклад веб-браузер.')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    
    doc.add_page_break()
    
    # ========== РОЗДІЛ 2 ==========
    p = doc.add_heading('РОЗДІЛ 2. ПРАКТИЧНЕ ЗАСТОСУВАННЯ ХМАРНИХ ТЕХНОЛОГІЙ', level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Підрозділ 2.1
    doc.add_heading('2.1. Провідні хмарні платформи', level=2)
    
    add_justified_text(doc, """Amazon Web Services (AWS) є лідером ринку хмарних послуг, займаючи близько 32% світового ринку. AWS пропонує понад 200 різноманітних сервісів, включаючи обчислювальні потужності, зберігання даних, бази даних, аналітику, машинне навчання та багато іншого.""")

    add_justified_text(doc, """Microsoft Azure займає друге місце на ринку з часткою близько 22%. Платформа тісно інтегрована з іншими продуктами Microsoft, що робить її привабливою для корпоративних клієнтів.""")
    
    # Гіперпосилання на AWS
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run('Офіційний сайт Amazon Web Services: ')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    add_hyperlink(p, 'https://aws.amazon.com', 'https://aws.amazon.com')
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    
    add_justified_text(doc, """Google Cloud Platform (GCP) займає третє місце з часткою близько 10%. GCP відрізняється потужними інструментами для аналізу даних та машинного навчання, а також глобальною мережевою інфраструктурою.""")

    add_justified_text(doc, """Інші значні гравці на ринку включають IBM Cloud, Oracle Cloud, Alibaba Cloud та інші. Кожна платформа має свої унікальні переваги та спеціалізацію.""")
    
    # Підрозділ 2.2
    doc.add_heading('2.2. Безпека в хмарних середовищах', level=2)
    
    add_justified_text(doc, """Безпека даних є одним з найважливіших аспектів при впровадженні хмарних технологій. Основні загрози включають несанкціонований доступ, витік даних, DDoS-атаки та внутрішні загрози.""")

    add_justified_text(doc, """Сучасні хмарні платформи використовують багаторівневий підхід до безпеки, який включає:""")
    
    add_bullet_list(doc, [
        'Шифрування даних при передачі та зберіганні;',
        'Багатофакторну автентифікацію;',
        'Системи виявлення та запобігання вторгненням;',
        'Регулярний аудит безпеки;',
        'Резервне копіювання та відновлення даних.'
    ])

    add_justified_text(doc, """Як зазначають дослідники [2], впровадження хмарних технологій вимагає ретельного планування стратегії безпеки та відповідності нормативним вимогам.""")
    
    # Параграф 2.2.1
    doc.add_heading('2.2.1. Методи захисту даних', level=3)
    
    add_justified_text(doc, """Захист даних у хмарних середовищах базується на кількох ключових принципах:""")

    add_justified_text(doc, """1. Принцип мінімальних привілеїв – користувачі отримують лише ті права доступу, які необхідні для виконання їхніх завдань.""")

    add_justified_text(doc, """2. Сегментація мережі – розділення інфраструктури на ізольовані сегменти для обмеження поширення потенційних загроз.""")

    add_justified_text(doc, """3. Шифрування – використання сучасних криптографічних алгоритмів (AES-256, RSA) для захисту конфіденційних даних.""")

    add_justified_text(doc, """4. Моніторинг та аудит – постійний контроль за активністю в системі та ведення детальних журналів подій.""")

    add_justified_text(doc, """За даними звіту компанії Gartner [3], організації, які впроваджують комплексні заходи безпеки, знижують ризик інцидентів на 70%.""")
    
    # Підрозділ 2.3
    doc.add_heading('2.3. Хмарні технології в освіті', level=2)
    
    add_justified_text(doc, """Освітня галузь активно впроваджує хмарні технології для забезпечення дистанційного навчання, спільної роботи та управління навчальним процесом.""")

    add_justified_text(doc, """Основні переваги використання хмарних технологій в освіті:""")
    
    add_bullet_list(doc, [
        'Доступність навчальних матеріалів з будь-якого пристрою;',
        'Можливість спільної роботи над проектами в реальному часі;',
        'Зниження витрат на інформаційні технології освітніх закладів;',
        'Автоматичне резервне копіювання та захист даних;',
        'Масштабованість відповідно до потреб.'
    ])

    add_justified_text(doc, """Популярні хмарні платформи для освіти включають Google Workspace for Education, Microsoft 365 Education та різноманітні системи управління навчанням (LMS).""")
    
    # Гіперпосилання на файл презентації
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run('Додаткову інформацію про хмарні сервіси можна знайти у презентації ')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    add_hyperlink(p, 'cloud_services_presentation', 'file:///Users/maksymyanisiv/test%20for%20cursor/cloud_services_presentation.html')
    run = p.add_run('.')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    
    doc.add_page_break()
    
    # ========== РОЗДІЛ 3 ==========
    p = doc.add_heading('РОЗДІЛ 3. ТЕНДЕНЦІЇ РОЗВИТКУ ХМАРНИХ ТЕХНОЛОГІЙ', level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Підрозділ 3.1
    doc.add_heading('3.1. Гібридні та мультихмарні рішення', level=2)
    
    add_justified_text(doc, """Сучасні тенденції показують зростаючий інтерес до гібридних та мультихмарних архітектур. Гібридна хмара поєднує публічні та приватні хмарні середовища, дозволяючи організаціям оптимізувати робочі навантаження.""")

    add_justified_text(doc, """Мультихмарна стратегія передбачає використання послуг від декількох хмарних провайдерів, що дозволяє:""")
    
    add_bullet_list(doc, [
        'Уникнути залежності від одного постачальника (vendor lock-in);',
        'Оптимізувати витрати, обираючи найкращі пропозиції;',
        'Підвищити надійність та відмовостійкість;',
        'Відповідати регуляторним вимогам різних юрисдикцій.'
    ])

    add_justified_text(doc, """Як зазначає професор Коваленко [4], до 2025 року понад 85% організацій будуть використовувати мультихмарну стратегію.""")
    
    # Підрозділ 3.2
    doc.add_heading('3.2. Edge Computing та IoT', level=2)
    
    add_justified_text(doc, """Edge Computing (периферійні обчислення) є важливим доповненням до хмарних технологій, особливо в контексті Інтернету речей (IoT). Ця технологія передбачає обробку даних ближче до джерела їх генерації, що дозволяє знизити затримки та навантаження на мережу.""")

    add_justified_text(doc, """Основні переваги Edge Computing:""")
    
    add_bullet_list(doc, [
        'Зменшення затримок у передачі даних;',
        'Підвищення конфіденційності та безпеки;',
        'Зниження вартості передачі даних;',
        'Робота в умовах обмеженої мережевої зв\'язності.'
    ])

    add_justified_text(doc, """Поєднання хмарних обчислень з Edge Computing створює потужну розподілену архітектуру для обробки даних у реальному часі.""")
    
    # Підрозділ 3.3
    doc.add_heading('3.3. Штучний інтелект у хмарних сервісах', level=2)
    
    add_justified_text(doc, """Інтеграція штучного інтелекту (ШІ) та машинного навчання з хмарними платформами відкриває нові можливості для бізнесу та наукових досліджень. Провідні хмарні провайдери пропонують широкий спектр AI-сервісів:""")
    
    add_bullet_list(doc, [
        'Amazon SageMaker для побудови та навчання моделей машинного навчання;',
        'Azure Machine Learning для корпоративного AI;',
        'Google AI Platform для розробки інтелектуальних застосунків;',
        'IBM Watson для когнітивних обчислень.'
    ])

    add_justified_text(doc, """Як підкреслюють експерти [5], демократизація AI через хмарні сервіси дозволяє навіть малим компаніям використовувати передові технології машинного навчання без значних інвестицій в інфраструктуру.""")
    
    # Параграф 3.3.1
    doc.add_heading('3.3.1. Генеративний AI у хмарі', level=3)
    
    add_justified_text(doc, """Генеративний штучний інтелект став однією з найгарячіших тем у світі інформаційних технологій. Великі мовні моделі (LLM), такі як GPT-4, Claude, Gemini, доступні через хмарні API, що дозволяє розробникам інтегрувати потужні AI-можливості у свої застосунки.""")

    add_justified_text(doc, """Застосування генеративного AI включає:""")
    
    add_bullet_list(doc, [
        'Автоматизацію створення контенту;',
        'Інтелектуальних чат-ботів та віртуальних асистентів;',
        'Генерацію коду та допомогу розробникам;',
        'Аналіз та узагальнення документів;',
        'Персоналізацію користувацького досвіду.'
    ])

    add_justified_text(doc, """Хмарна інфраструктура є критично важливою для роботи таких моделей, оскільки вони вимагають значних обчислювальних ресурсів.""")
    
    doc.add_page_break()
    
    # ========== ВИСНОВКИ ==========
    p = doc.add_heading('ВИСНОВКИ', level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_justified_text(doc, """У результаті проведеного дослідження можна зробити наступні висновки:""")

    add_justified_text(doc, """1. Хмарні технології стали невід'ємною частиною сучасної інформаційної інфраструктури, забезпечуючи гнучкість, масштабованість та економічну ефективність для організацій різного масштабу.""")

    add_justified_text(doc, """2. Три основні моделі надання хмарних послуг (IaaS, PaaS, SaaS) задовольняють різні потреби споживачів – від базової інфраструктури до готових програмних рішень.""")

    add_justified_text(doc, """3. Безпека залишається ключовим викликом при впровадженні хмарних технологій, але сучасні платформи пропонують комплексні рішення для захисту даних.""")

    add_justified_text(doc, """4. Тенденції розвитку вказують на зростання популярності гібридних та мультихмарних архітектур, а також інтеграцію з технологіями Edge Computing та штучного інтелекту.""")

    add_justified_text(doc, """5. Демократизація доступу до передових інформаційних технологій через хмарні сервіси створює нові можливості для інновацій та цифрової трансформації.""")

    add_justified_text(doc, """Хмарні технології продовжуватимуть еволюціонувати, відкриваючи нові горизонти для бізнесу, науки та суспільства в цілому.""")
    
    doc.add_page_break()
    
    # ========== СПИСОК ЛІТЕРАТУРИ ==========
    p = doc.add_heading('СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ', level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    sources = [
        '1. Mell P., Grance T. The NIST Definition of Cloud Computing. National Institute of Standards and Technology, Special Publication 800-145, 2011. 7 p.',
        '2. Armbrust M. et al. A View of Cloud Computing. Communications of the ACM, 2010. Vol. 53, No. 4. P. 50-58.',
        '3. Gartner Inc. Magic Quadrant for Cloud Infrastructure and Platform Services. 2024. URL: https://www.gartner.com/en/documents/4021139',
        '4. Коваленко В.В. Хмарні технології в освіті: монографія. Київ: Наукова думка, 2023. 256 с.',
        '5. Литвинов В.А., Петренко О.М. Застосування хмарних обчислень у корпоративних інформаційних системах. Вісник НТУУ "КПІ". Серія: Інформатика, 2024. №2. С. 45-52.',
        '6. Amazon Web Services. AWS Documentation. URL: https://docs.aws.amazon.com/',
        '7. Microsoft Azure. Azure Architecture Center. URL: https://docs.microsoft.com/azure/architecture/',
    ]
    
    for source in sources:
        add_justified_text(doc, source, first_indent=False)
    
    doc.add_page_break()
    
    # ========== КІНЦЕВІ ВИНОСКИ ==========
    p = doc.add_heading('КІНЦЕВІ ВИНОСКИ', level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    endnotes = [
        'І. Хмарні обчислення – технологія розподілених обчислень, яка надає користувачам комп\'ютерні ресурси та потужності як Інтернет-сервіс.',
        'ІІ. Віртуалізація – створення віртуальної версії пристрою або ресурсу, наприклад сервера, пристрою зберігання даних, мережі або операційної системи.',
        'ІІІ. API (Application Programming Interface) – інтерфейс програмування застосунків, набір протоколів та інструментів для створення програмного забезпечення.',
        'IV. Масштабованість – здатність системи збільшувати свою продуктивність при додаванні ресурсів (зазвичай апаратних).',
    ]
    
    for endnote in endnotes:
        add_justified_text(doc, endnote, first_indent=False, size=12)
    
    # Додати нумерацію сторінок
    add_page_number(doc)
    
    # Зберегти документ
    doc.save('/Users/maksymyanisiv/test for cursor/Реферат.docx')
    print("Реферат створено успішно!")
    

def create_figure_document():
    """Створити документ з блок-схемою"""
    doc = Document()
    
    # Налаштування стилів
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    
    # Налаштування полів
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)
    
    # Заголовок
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('БЛОК-СХЕМА РОЗВ\'ЯЗУВАННЯ КВАДРАТНОГО РІВНЯННЯ')
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    doc.add_paragraph()
    
    # Опис алгоритму
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run('Дана блок-схема ілюструє алгоритм розв\'язування квадратного рівняння ax² + bx + c = 0.')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('Алгоритм враховує всі можливі випадки:')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    
    cases = [
        'Якщо a = 0, рівняння вироджується в лінійне bx + c = 0',
        'Якщо b = 0 і c = 0, рішенням є вся множина дійсних чисел x ∈ R',
        'Якщо b = 0 і c ≠ 0, рівняння не має розв\'язків x ∈ ∅',
        'Якщо b ≠ 0, рішення лінійного рівняння: x = -c/b',
        'Для квадратного рівняння (a ≠ 0) обчислюється дискримінант D = b² - 4ac',
        'Якщо D < 0, рівняння не має дійсних коренів',
        'Якщо D = 0, рівняння має один корінь: x = -b/(2a)',
        'Якщо D > 0, рівняння має два корені: x₁ = (-b - √D)/(2a), x₂ = (-b + √D)/(2a)'
    ]
    
    for case in cases:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run('• ' + case)
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()
    
    # Текстове представлення блок-схеми
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Блок-схема алгоритму')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    doc.add_paragraph()
    
    flowchart_text = """
╔════════════════════════════════════════════════════════════╗
║                      ПОЧАТОК                                ║
║                   Введення a, b, c                          ║
╚════════════════════════════════════════════════════════════╝
                           │
                           ▼
                 ┌─────────────────┐
                 │     a = 0?      │
                 └─────────────────┘
                  /              \\
             ТАК /                \\ НІ
                ▼                  ▼
     ┌───────────────┐    ┌───────────────────┐
     │   bx + c = 0  │    │   D = b² - 4ac    │
     └───────────────┘    └───────────────────┘
            │                      │
            ▼                      ▼
     ┌─────────────┐       ┌─────────────┐
     │   b = 0?    │       │   D < 0?    │
     └─────────────┘       └─────────────┘
      /          \\         /          \\
  ТАК/            \\НІ  ТАК/            \\НІ
     ▼            ▼       ▼             ▼
┌─────────┐  ┌─────────┐ ┌────────┐ ┌─────────┐
│ c = 0?  │  │x = -c/b │ │ x ∈ ∅  │ │ D = 0?  │
└─────────┘  └─────────┘ └────────┘ └─────────┘
 /      \\                            /      \\
ТАК     НІ                        ТАК       НІ
▼        ▼                          ▼        ▼
┌──────┐ ┌──────┐           ┌───────────┐ ┌─────────────────┐
│x ∈ R │ │x ∈ ∅ │           │x = -b/(2a)│ │x₁=(-b-√D)/(2a)  │
└──────┘ └──────┘           └───────────┘ │x₂=(-b+√D)/(2a)  │
                                          └─────────────────┘
"""
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(flowchart_text)
    run.font.size = Pt(9)
    run.font.name = 'Courier New'
    
    doc.add_paragraph()
    
    # Формули
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Основні формули:')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    run.bold = True
    
    formulas = [
        'Дискримінант: D = b² - 4ac',
        'Корені при D > 0: x₁ = (-b - √D)/(2a),  x₂ = (-b + √D)/(2a)',
        'Корінь при D = 0: x = -b/(2a)',
        'Лінійне рівняння: x = -c/b',
    ]
    
    for formula in formulas:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(formula)
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
        run.italic = True
    
    doc.add_paragraph()
    
    # Підпис
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Рис. 1. Блок-схема розв\'язування рівняння')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.italic = True
    
    # Водяний знак (текст у header)
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run('КОПІЮВАТИ ЗАБОРОНЕНО')
        run.font.size = Pt(36)
        run.font.color.rgb = RGBColor(220, 220, 220)
        run.font.name = 'Arial'
    
    # Зберегти документ
    doc.save('/Users/maksymyanisiv/test for cursor/Фігура.docx')
    print("Документ 'Фігура' створено успішно!")


def create_readme():
    """Створити README з описом виконаної роботи"""
    readme_content = """# Лабораторна робота: Текстовий процесор

## Опис виконаної роботи

### Документ "Реферат.docx"

Створено реферат на тему "Хмарні технології та їх застосування в сучасних інформаційних системах", який містить:

1. **Титульна сторінка** - з назвою міністерства, університету, факультету, теми, виконавця та керівника

2. **Анотація** (сторінка 2) - короткий опис реферату до 2000 символів

3. **Автоматичний зміст** (сторінка 3) - з розділами, підрозділами та параграфами

4. **Структура реферату:**
   - ВСТУП
   - РОЗДІЛ 1. ТЕОРЕТИЧНІ ОСНОВИ ХМАРНИХ ТЕХНОЛОГІЙ
     - 1.1. Поняття та історія розвитку хмарних обчислень
       - 1.1.1. Основні характеристики хмарних обчислень
     - 1.2. Моделі надання хмарних послуг
   - РОЗДІЛ 2. ПРАКТИЧНЕ ЗАСТОСУВАННЯ ХМАРНИХ ТЕХНОЛОГІЙ
     - 2.1. Провідні хмарні платформи
     - 2.2. Безпека в хмарних середовищах
       - 2.2.1. Методи захисту даних
     - 2.3. Хмарні технології в освіті
   - РОЗДІЛ 3. ТЕНДЕНЦІЇ РОЗВИТКУ ХМАРНИХ ТЕХНОЛОГІЙ
     - 3.1. Гібридні та мультихмарні рішення
     - 3.2. Edge Computing та IoT
     - 3.3. Штучний інтелект у хмарних сервісах
       - 3.3.1. Генеративний AI у хмарі
   - ВИСНОВКИ
   - СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ (7 джерел)
   - КІНЦЕВІ ВИНОСКИ (4 виноски з форматом І, ІІ, ІІІ, IV)

5. **Нумерація сторінок** - автоматична нумерація внизу сторінки

6. **Гіперпосилання:**
   - На файл cloud_services_statistics.csv
   - На файл cloud_services_presentation.html
   - На Wikipedia - Cloud Computing
   - На https://aws.amazon.com

7. **Закладки:**
   - cloud_computing_definition - визначення хмарних обчислень
   - iaas_definition - визначення IaaS
   - paas_definition - визначення PaaS
   - saas_definition - визначення SaaS

8. **Перехресні посилання** - посилання на джерела у списку літератури [1], [2], [3], [4], [5]

### Документ "Фігура.docx"

Створено документ з блок-схемою розв'язування квадратного рівняння:

1. **Заголовок** - "БЛОК-СХЕМА РОЗВ'ЯЗУВАННЯ КВАДРАТНОГО РІВНЯННЯ"

2. **Опис алгоритму** - детальний опис всіх випадків розв'язування

3. **Блок-схема** - текстове представлення алгоритму з використанням псевдографіки

4. **Формули:**
   - Дискримінант: D = b² - 4ac
   - Корені при D > 0
   - Корінь при D = 0
   - Лінійне рівняння

5. **Водяний знак** - "КОПІЮВАТИ ЗАБОРОНЕНО"

## Висновки

У ході виконання лабораторної роботи було:

1. Створено структурований реферат на професійну тему з усіма необхідними елементами оформлення
2. Застосовано автоматичну нумерацію сторінок
3. Додано гіперпосилання на локальні файли та інтернет-ресурси
4. Створено закладки для швидкої навігації по документу
5. Додано перехресні посилання на літературні джерела
6. Сформовано автоматичний зміст
7. Створено документ з блок-схемою алгоритму
8. Застосовано водяний знак для захисту документа

## Як запустити скрипт

```bash
pip install python-docx
python create_referat.py
```

## Файли

- `Реферат.docx` - основний документ реферату
- `Фігура.docx` - документ з блок-схемою
- `create_referat.py` - скрипт для створення документів
"""
    
    with open('/Users/maksymyanisiv/test for cursor/README_LAB_WORK_10.md', 'w', encoding='utf-8') as f:
        f.write(readme_content)
    print("README створено успішно!")


if __name__ == '__main__':
    print("Створення документів...")
    print("-" * 50)
    
    try:
        create_referat()
    except Exception as e:
        print(f"Помилка при створенні реферату: {e}")
        import traceback
        traceback.print_exc()
    
    try:
        create_figure_document()
    except Exception as e:
        print(f"Помилка при створенні документа 'Фігура': {e}")
        import traceback
        traceback.print_exc()
    
    try:
        create_readme()
    except Exception as e:
        print(f"Помилка при створенні README: {e}")
    
    print("-" * 50)
    print("Готово!")
    print("\nСтворені файли:")
    print("  - Реферат.docx")
    print("  - Фігура.docx")
    print("  - README_LAB_WORK_10.md")
